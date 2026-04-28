#!/usr/bin/env python3
"""
Generate FGA Test Matrix as a Word document (.docx)
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_test_matrix_docx():
    doc = Document()

    # Title
    title = doc.add_heading('FGA Test Matrix - ProGear Sales AI', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('FGA Model: ProGear New (01KQAYNTB7AY897GMG5EHDVM8K)')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x01, 0x31, 0xFF)

    doc.add_paragraph()

    # Item Requirements Section
    doc.add_heading('Item Clearance Requirements', level=1)
    item_table = doc.add_table(rows=3, cols=2)
    item_table.style = 'Table Grid'

    # Header
    hdr = item_table.rows[0].cells
    hdr[0].text = 'Item'
    hdr[1].text = 'Required Clearance'
    for cell in hdr:
        set_cell_shading(cell, '00033E')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True

    # Data
    items = [
        ('widget-a (basketballs, general inventory)', '3'),
        ('classified-part (sensitive military items)', '7'),
    ]
    for i, (item, clearance) in enumerate(items, 1):
        row = item_table.rows[i].cells
        row[0].text = item
        row[1].text = clearance

    doc.add_paragraph()

    # Test Scenarios Section
    doc.add_heading('Test Scenarios', level=1)

    # Define test data with updated prompts using actual item names
    scenarios = [
        # Section header
        ('section', 'Manager Scenarios', '', '', '', '', '', '', ''),
        # Manager scenarios
        ('1', 'Manager view widget-a', 'Bob', 'true', 'false', '5',
         'How many widget-a items are in stock?', 'PASS', 'active_manager can view'),
        ('2', 'Manager update widget-a (sufficient clearance)', 'Bob', 'true', 'false', '5',
         'Add 50 widget-a to inventory', 'PASS', 'Clearance 5 ≥ required 3'),
        ('3', 'Manager update classified-part (insufficient)', 'Bob', 'true', 'false', '2',
         'Update classified-part quantity to 100', 'FAIL', 'Clearance 2 < required 7'),
        ('4', 'Manager view classified-part (low clearance)', 'Bob', 'true', 'false', '2',
         'Check stock for classified-part', 'PASS', 'View doesn\'t need clearance'),
        ('5', 'Manager on vacation (view blocked)', 'Bob', 'true', 'true', '5',
         'How many widget-a items are in stock?', 'FAIL', 'on_vacation blocks active_manager'),
        ('6', 'Manager on vacation (update blocked)', 'Bob', 'true', 'true', '10',
         'Add 50 widget-a to inventory', 'FAIL', 'Vacation blocks even with max clearance'),
        # Section header
        ('section', 'Viewer Scenarios (NEW)', '', '', '', '', '', '', ''),
        # Viewer scenarios
        ('7', 'Non-manager view widget-a (viewer role)', 'Sarah', 'false', 'false', '7',
         'How many widget-a items are in stock?', 'PASS', 'active_viewer can view'),
        ('8', 'Non-manager update widget-a (denied)', 'Sarah', 'false', 'false', '7',
         'Add 50 widget-a to inventory', 'FAIL', 'Viewers cannot update'),
        ('9', 'Non-manager on vacation', 'Sarah', 'false', 'true', '7',
         'How many widget-a items are in stock?', 'FAIL', 'on_vacation blocks active_viewer'),
        # Section header
        ('section', 'Update Clearance Scenarios', '', '', '', '', '', '', ''),
        # Clearance scenarios
        ('10', 'Manager update classified-part (sufficient)', 'Bob', 'true', 'false', '8',
         'Update classified-part quantity to 50', 'PASS', 'Clearance 8 ≥ required 7'),
        ('11', 'Manager update classified-part (exact match)', 'Bob', 'true', 'false', '7',
         'Update classified-part quantity to 50', 'PASS', 'Clearance 7 = required 7'),
        ('12', 'Viewer update widget-a (denied)', 'Sarah', 'false', 'false', '10',
         'Add 100 widget-a to inventory', 'FAIL', 'Viewers can\'t update even with clearance 10'),
    ]

    # Create main table
    table = doc.add_table(rows=1, cols=9)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ['#', 'Scenario', 'User', 'Manager', 'Vacation', 'Clearance', 'Prompt', 'Expected', 'Key Point']
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], '00033E')
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(9)

    # Data rows
    for row_data in scenarios:
        row_cells = table.add_row().cells

        if row_data[0] == 'section':
            # Section header row
            row_cells[0].merge(row_cells[8])
            row_cells[0].text = row_data[1]
            set_cell_shading(row_cells[0], 'E0DAD2')
            for paragraph in row_cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
        else:
            # Data row
            for i, value in enumerate(row_data):
                row_cells[i].text = value
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)
                        # Color coding
                        if value == 'PASS':
                            run.font.color.rgb = RGBColor(0x22, 0xc5, 0x5e)
                            run.font.bold = True
                        elif value == 'FAIL':
                            run.font.color.rgb = RGBColor(0xef, 0x44, 0x44)
                            run.font.bold = True
                        elif value == 'true':
                            run.font.color.rgb = RGBColor(0x01, 0x31, 0xFF)
                        elif value == 'false':
                            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(0.3)   # #
        row.cells[1].width = Inches(1.8)   # Scenario
        row.cells[2].width = Inches(0.5)   # User
        row.cells[3].width = Inches(0.5)   # Manager
        row.cells[4].width = Inches(0.5)   # Vacation
        row.cells[5].width = Inches(0.6)   # Clearance
        row.cells[6].width = Inches(2.5)   # Prompt
        row.cells[7].width = Inches(0.6)   # Expected
        row.cells[8].width = Inches(1.8)   # Key Point

    doc.add_paragraph()

    # Permission Matrix Section
    doc.add_heading('Permission Matrix', level=1)

    perm_table = doc.add_table(rows=7, cols=4)
    perm_table.style = 'Table Grid'

    perm_headers = ['User State', 'can_view', 'can_update (widget-a)', 'can_update (classified-part)']
    perm_hdr = perm_table.rows[0].cells
    for i, header in enumerate(perm_headers):
        perm_hdr[i].text = header
        set_cell_shading(perm_hdr[i], '00033E')
        for paragraph in perm_hdr[i].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(9)

    perm_data = [
        ('Manager=true, Vacation=false, CL=5', '✓', '✓', '✗'),
        ('Manager=true, Vacation=false, CL=7', '✓', '✓', '✓'),
        ('Manager=true, Vacation=false, CL=2', '✓', '✗', '✗'),
        ('Manager=true, Vacation=true, CL=10', '✗', '✗', '✗'),
        ('Manager=false, Vacation=false, CL=7 (Viewer)', '✓', '✗', '✗'),
        ('Manager=false, Vacation=true, CL=7', '✗', '✗', '✗'),
    ]

    for i, row_data in enumerate(perm_data, 1):
        row_cells = perm_table.rows[i].cells
        for j, value in enumerate(row_data):
            row_cells[j].text = value
            for paragraph in row_cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    if value == '✓':
                        run.font.color.rgb = RGBColor(0x22, 0xc5, 0x5e)
                        run.font.bold = True
                    elif value == '✗':
                        run.font.color.rgb = RGBColor(0xef, 0x44, 0x44)
                        run.font.bold = True

    doc.add_paragraph()

    # Legend
    doc.add_heading('Legend', level=2)
    legend = doc.add_paragraph()
    legend.add_run('CL = Clearance Level\n').font.size = Pt(10)
    legend.add_run('✓ = Allowed, ✗ = Denied\n').font.size = Pt(10)
    pass_run = legend.add_run('PASS')
    pass_run.font.color.rgb = RGBColor(0x22, 0xc5, 0x5e)
    pass_run.font.bold = True
    legend.add_run(' = Access granted\n').font.size = Pt(10)
    fail_run = legend.add_run('FAIL')
    fail_run.font.color.rgb = RGBColor(0xef, 0x44, 0x44)
    fail_run.font.bold = True
    legend.add_run(' = Access denied').font.size = Pt(10)

    doc.add_paragraph()

    # FGA Model Section
    doc.add_heading('FGA Model Relations', level=1)
    model_code = """type inventory_system
  relations
    define viewer: [user]                              # NEW
    define manager: [user]
    define on_vacation: [user]
    define active_manager: manager but not on_vacation
    define active_viewer: viewer but not on_vacation   # NEW
    define can_manage: active_manager
    define can_read: active_manager or active_viewer   # NEW

type inventory_item
  relations
    define can_view: can_read from parent              # CHANGED
    define can_update: has_clearance and can_manage from parent"""

    code_para = doc.add_paragraph()
    code_run = code_para.add_run(model_code)
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(9)

    # Save
    output_path = '/Users/rajeshkumar/Documents/GitHub/courtedge-ai-demo/docs/FGA_TEST_MATRIX_COMPLETE.docx'
    doc.save(output_path)
    print(f'Document saved to: {output_path}')
    return output_path

if __name__ == '__main__':
    create_test_matrix_docx()
